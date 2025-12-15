"""Resume parsing and management service."""

import asyncio
import io
import logging
import re
from concurrent.futures import ThreadPoolExecutor
from datetime import date, datetime
from typing import Optional, Dict, Any, List, BinaryIO
from uuid import UUID

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from backend.models.resume import Resume, WorkExperience
from backend.services.llm_service import get_llm_service, LLMService

logger = logging.getLogger(__name__)

# Placeholder patterns that LLMs commonly generate - these should be rejected
PLACEHOLDER_PATTERNS = {
    "emails": [
        "example.com", "test.com", "email.com", "mail.com", "sample.com",
        "placeholder", "your-email", "youremail", "user@", "name@",
    ],
    "companies": [
        "techcorp", "acme", "company x", "company y", "company z",
        "abc corp", "xyz corp", "sample company", "example company",
        "tech company", "startup x", "corporation", "company name",
        "previous company", "current company", "your company",
    ],
    "names": [
        "john doe", "jane doe", "john smith", "jane smith",
        "your name", "full name", "first last", "name here",
    ],
}


def _is_placeholder(value: Optional[str], category: str) -> bool:
    """Check if a value appears to be a placeholder."""
    if not value:
        return False
    value_lower = value.lower().strip()
    patterns = PLACEHOLDER_PATTERNS.get(category, [])
    for pattern in patterns:
        if pattern in value_lower:
            return True
    return False


def _clean_llm_artifacts(value: Optional[str]) -> Optional[str]:
    """Clean up common LLM artifacts from extracted values.

    LLMs sometimes add prefixes like '#', 'I ', '* ', etc. to values.
    This function strips those artifacts to get the clean value.
    """
    if not value:
        return value

    original = value

    # Strip whitespace first
    value = value.strip()

    # Remove common LLM prefix artifacts
    # Pattern: starts with # or * or - or | followed by optional space
    value = re.sub(r'^[#\*\-\|]\s*', '', value)

    # Pattern: starts with a single letter followed by space (like "I 860...")
    # Only if followed by a digit or @ (for phone/email)
    value = re.sub(r'^[A-Za-z]\s+(?=[\d@])', '', value)

    # Remove markdown-style formatting
    value = re.sub(r'^\*\*(.+)\*\*$', r'\1', value)  # **bold**
    value = re.sub(r'^__(.+)__$', r'\1', value)  # __underline__

    # Strip any remaining whitespace
    value = value.strip()

    if value != original:
        logger.debug(f"Cleaned LLM artifact: '{original}' -> '{value}'")

    return value if value else None


def _clean_email(value: Optional[str]) -> Optional[str]:
    """Extract and clean email from potentially messy LLM output."""
    if not value:
        return None

    # First apply general cleaning
    value = _clean_llm_artifacts(value) or ""

    # Try to extract email pattern from string
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, value)
    if match:
        return match.group(0).lower()

    return None


def _clean_phone(value: Optional[str]) -> Optional[str]:
    """Extract and clean phone number from potentially messy LLM output."""
    if not value:
        return None

    # First apply general cleaning
    value = _clean_llm_artifacts(value) or ""

    # Extract just digits, dots, dashes, parens, plus, and spaces
    # Keep formatting for readability
    cleaned = re.sub(r'[^\d\.\-\(\)\+\s]', '', value)
    cleaned = cleaned.strip()

    # Must have at least 10 digits to be valid
    digits_only = re.sub(r'\D', '', cleaned)
    if len(digits_only) >= 10:
        return cleaned

    return None


def _sanitize_parsed_data(data: Dict[str, Any]) -> Dict[str, Any]:
    """Remove placeholder values and clean LLM artifacts from parsed resume data."""

    # Clean and sanitize email
    raw_email = data.get("email")
    cleaned_email = _clean_email(raw_email)
    if cleaned_email:
        if _is_placeholder(cleaned_email, "emails"):
            logger.warning(f"Removing placeholder email: {cleaned_email}")
            data["email"] = None
        else:
            data["email"] = cleaned_email
            if raw_email != cleaned_email:
                logger.info(f"Cleaned email: '{raw_email}' -> '{cleaned_email}'")
    else:
        data["email"] = None

    # Clean and sanitize phone
    raw_phone = data.get("phone")
    cleaned_phone = _clean_phone(raw_phone)
    if cleaned_phone:
        data["phone"] = cleaned_phone
        if raw_phone != cleaned_phone:
            logger.info(f"Cleaned phone: '{raw_phone}' -> '{cleaned_phone}'")
    else:
        data["phone"] = None

    # Clean and sanitize name
    raw_name = data.get("full_name")
    cleaned_name = _clean_llm_artifacts(raw_name)
    if cleaned_name:
        if _is_placeholder(cleaned_name, "names"):
            logger.warning(f"Removing placeholder name: {cleaned_name}")
            data["full_name"] = None
        else:
            data["full_name"] = cleaned_name
    else:
        data["full_name"] = None

    # Sanitize work experiences
    work_experiences = data.get("work_experiences", [])
    sanitized_experiences = []
    for exp in work_experiences:
        company = _clean_llm_artifacts(exp.get("company", ""))
        if not company or _is_placeholder(company, "companies"):
            logger.warning(f"Removing work experience with placeholder company: {exp.get('company')}")
            continue  # Skip this experience entirely
        exp["company"] = company
        # Also clean title and location
        if exp.get("title"):
            exp["title"] = _clean_llm_artifacts(exp["title"])
        if exp.get("location"):
            exp["location"] = _clean_llm_artifacts(exp["location"])
        sanitized_experiences.append(exp)
    data["work_experiences"] = sanitized_experiences

    return data


class ResumeService:
    """Service for parsing and managing resumes."""

    def __init__(self, db: AsyncSession, llm_service: Optional[LLMService] = None):
        """Initialize resume service.

        Args:
            db: Database session
            llm_service: Optional LLM service (defaults to global instance)
        """
        self.db = db
        self.llm = llm_service or get_llm_service()

    async def upload_and_parse(
        self,
        user_id: UUID,
        file_content: bytes,
        file_name: str,
        file_type: str,
    ) -> Resume:
        """Upload and parse a resume file.

        Args:
            user_id: User ID to associate resume with
            file_content: Raw file bytes
            file_name: Original file name
            file_type: MIME type or extension (pdf, docx, txt)

        Returns:
            Parsed Resume object
        """
        logger.info(f"Processing resume for user {user_id}: {file_name}")

        # Extract text from file
        raw_text = await self._extract_text(file_content, file_type)
        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError("Could not extract meaningful text from resume")

        # Log extracted text for debugging
        logger.info(f"EXTRACTED TEXT LENGTH: {len(raw_text)} chars")
        logger.info(f"EXTRACTED TEXT PREVIEW (first 500 chars):\n{raw_text[:500]}")
        logger.info(f"EXTRACTED TEXT PREVIEW (last 300 chars):\n{raw_text[-300:]}")

        # Parse with LLM
        parsed_data = await self._parse_with_llm(raw_text)

        # Check if user already has a resume
        existing = await self.get_resume(user_id)
        if existing:
            # Update existing resume
            resume = existing
            resume.file_name = file_name
            resume.file_type = self._normalize_file_type(file_type)
            resume.file_size = len(file_content)
            resume.raw_text = raw_text
            resume.parsed_at = datetime.utcnow()

            # Clear old work experiences
            resume.work_experiences = []
        else:
            # Create new resume
            resume = Resume(
                user_id=user_id,
                file_name=file_name,
                file_type=self._normalize_file_type(file_type),
                file_size=len(file_content),
                raw_text=raw_text,
                parsed_at=datetime.utcnow(),
            )
            self.db.add(resume)

        # Apply parsed data
        self._apply_parsed_data(resume, parsed_data)

        await self.db.flush()
        await self.db.refresh(resume)

        logger.info(f"Resume parsed successfully for user {user_id}")
        return resume

    async def parse_text(self, user_id: UUID, text: str) -> Resume:
        """Parse resume from pasted text.

        Args:
            user_id: User ID
            text: Resume text content

        Returns:
            Parsed Resume object
        """
        logger.info(f"Parsing resume text for user {user_id}")

        if len(text.strip()) < 50:
            raise ValueError("Resume text is too short")

        # Parse with LLM
        parsed_data = await self._parse_with_llm(text)

        # Check if user already has a resume
        existing = await self.get_resume(user_id)
        if existing:
            resume = existing
            resume.raw_text = text
            resume.file_name = None
            resume.file_type = "text"
            resume.parsed_at = datetime.utcnow()
            resume.work_experiences = []
        else:
            resume = Resume(
                user_id=user_id,
                raw_text=text,
                file_type="text",
                parsed_at=datetime.utcnow(),
            )
            self.db.add(resume)

        self._apply_parsed_data(resume, parsed_data)

        await self.db.flush()
        await self.db.refresh(resume)

        return resume

    async def get_resume(self, user_id: UUID) -> Optional[Resume]:
        """Get resume for a user.

        Args:
            user_id: User ID

        Returns:
            Resume if found, None otherwise
        """
        result = await self.db.execute(
            select(Resume)
            .where(Resume.user_id == user_id)
            .options(selectinload(Resume.work_experiences))
        )
        return result.scalar_one_or_none()

    async def delete_resume(self, user_id: UUID) -> bool:
        """Delete a user's resume.

        Args:
            user_id: User ID

        Returns:
            True if deleted, False if not found
        """
        resume = await self.get_resume(user_id)
        if resume:
            await self.db.delete(resume)
            return True
        return False

    async def get_relevant_experience(
        self,
        user_id: UUID,
        required_skills: List[str],
        nice_to_have_skills: Optional[List[str]] = None,
        limit: int = 3,
    ) -> List[WorkExperience]:
        """Get work experiences most relevant to job requirements.

        Args:
            user_id: User ID
            required_skills: List of required skills from job
            nice_to_have_skills: Optional list of nice-to-have skills
            limit: Maximum number of experiences to return

        Returns:
            List of most relevant WorkExperience objects
        """
        resume = await self.get_resume(user_id)
        if not resume or not resume.work_experiences:
            return []

        # Score each experience by skill overlap
        all_target_skills = set(s.lower() for s in required_skills)
        if nice_to_have_skills:
            all_target_skills.update(s.lower() for s in nice_to_have_skills)

        scored_experiences = []
        for exp in resume.work_experiences:
            exp_skills = set(s.lower() for s in (exp.skills_used or []))
            # Also check description and achievements for skill mentions
            text = f"{exp.description or ''} {' '.join(exp.achievements or [])}".lower()

            skill_matches = sum(
                1 for skill in all_target_skills
                if skill in exp_skills or skill in text
            )

            # Bonus for required vs nice-to-have
            required_matches = sum(
                1 for skill in required_skills
                if skill.lower() in exp_skills or skill.lower() in text
            )

            score = skill_matches + (required_matches * 0.5)
            scored_experiences.append((exp, score))

        # Sort by score and return top N
        scored_experiences.sort(key=lambda x: x[1], reverse=True)
        return [exp for exp, _ in scored_experiences[:limit]]

    async def _extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text from file content.

        Args:
            file_content: Raw file bytes
            file_type: File type (pdf, docx, txt, etc.)

        Returns:
            Extracted text
        """
        file_type = self._normalize_file_type(file_type)

        if file_type == "pdf":
            return await self._extract_pdf_text(file_content)
        elif file_type == "docx":
            return await self._extract_docx_text(file_content)
        elif file_type in ("txt", "text"):
            return file_content.decode("utf-8", errors="ignore")
        else:
            # Try to decode as text
            return file_content.decode("utf-8", errors="ignore")

    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_pdf_text_sync, file_content)

    def _extract_pdf_text_sync(self, file_content: bytes) -> str:
        """Synchronous PDF text extraction (runs in thread pool)."""
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("pdfplumber not installed. Run: pip install pdfplumber")
            raise
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Could not extract text from PDF: {e}")

    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_docx_text_sync, file_content)

    def _extract_docx_text_sync(self, file_content: bytes) -> str:
        """Synchronous DOCX text extraction (runs in thread pool)."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_content))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Could not extract text from DOCX: {e}")

    def _normalize_file_type(self, file_type: str) -> str:
        """Normalize file type to simple extension."""
        file_type = file_type.lower()

        if "pdf" in file_type:
            return "pdf"
        elif "docx" in file_type or "document" in file_type:
            return "docx"
        elif "doc" in file_type:
            return "doc"
        elif "text" in file_type or "plain" in file_type:
            return "txt"
        else:
            return file_type

    async def _parse_with_llm(self, text: str) -> Dict[str, Any]:
        """Parse resume text using LLM.

        Args:
            text: Resume text content

        Returns:
            Parsed resume data as dictionary
        """
        system_prompt = """You are an expert resume parser. Your job is to extract EXACT information from resumes.

CRITICAL RULES:
1. Extract ONLY information that ACTUALLY EXISTS in the resume text
2. NEVER use placeholder values like "example.com", "TechCorp", "Acme Inc", or "John Doe"
3. If information is not present, use null - do NOT make up or guess values
4. Extract ALL work experiences - do not skip any positions
5. Use the EXACT company names, job titles, and email addresses as written in the resume
6. For dates, use ISO format (YYYY-MM-DD) or just year (YYYY) if month is not specified
7. For metrics and achievements, preserve exact numbers and percentages"""

        # Use more text to avoid truncating important information
        resume_text = text[:15000] if len(text) > 15000 else text

        prompt = f"""Parse this resume and extract structured information.

IMPORTANT: Extract EXACT values from the text. Do NOT use placeholder values.
If you cannot find a piece of information, use null instead of making something up.
Extract ALL work experiences listed - there may be many positions.

RESUME TEXT:
{resume_text}

Extract the following information in JSON format. Return ONLY valid JSON, no comments or explanations:

{{
    "full_name": "EXACT name from resume or null if not found",
    "email": "EXACT email address from resume or null if not found",
    "phone": "EXACT phone number from resume or null if not found",
    "location": "EXACT location from resume or null if not found",
    "linkedin_url": "EXACT LinkedIn URL from resume or null",
    "github_url": "EXACT GitHub URL from resume or null",
    "portfolio_url": "EXACT portfolio URL from resume or null",
    "summary": "professional summary/objective as written or null",
    "skills": ["extract ALL skills and technologies mentioned"],
    "certifications": ["extract ALL certifications mentioned"],
    "languages": ["extract spoken languages if mentioned"],
    "education": [
        {{
            "degree": "EXACT degree name",
            "field": "EXACT field of study",
            "school": "EXACT institution name",
            "year": "graduation year",
            "gpa": "GPA if mentioned or null"
        }}
    ],
    "work_experiences": [
        {{
            "company": "EXACT company name as written in resume",
            "title": "EXACT job title as written in resume",
            "location": "EXACT location or null",
            "employment_type": "full-time or contract or freelance or part-time or null",
            "is_remote": true or false,
            "start_date": "YYYY-MM-DD or YYYY-MM or YYYY",
            "end_date": "YYYY-MM-DD or YYYY-MM or YYYY or null if current position",
            "is_current": true or false,
            "description": "role description if provided",
            "achievements": ["ALL bullet points and achievements for this role"],
            "skills_used": ["technologies and skills used in this role"],
            "metrics": {{"extracted metrics like revenue, team size, etc"}}
        }}
    ],
    "parse_quality_score": 85
}}

REMINDERS:
- Include EVERY work experience, even if there are 5+ positions
- Use EXACT company names - never substitute with generic names
- The email should be the ACTUAL email from the resume, not a placeholder
- parse_quality_score should be 0-100 indicating extraction confidence"""

        try:
            result = await self.llm.generate_structured(prompt, system_prompt)

            # Log raw LLM response for debugging
            logger.info(f"RAW LLM Response - email: {result.get('email')}, name: {result.get('full_name')}")
            logger.info(f"RAW LLM Response - work_experiences count: {len(result.get('work_experiences', []))}")
            for i, exp in enumerate(result.get('work_experiences', [])[:3]):
                logger.info(f"  Experience {i+1}: {exp.get('company')} - {exp.get('title')}")

            # Remove placeholder values that LLMs sometimes generate
            result = _sanitize_parsed_data(result)
            logger.info(f"AFTER SANITIZATION - {len(result.get('work_experiences', []))} work experiences, email={result.get('email')}")
            return result
        except Exception as e:
            logger.error(f"LLM parsing failed: {e}")
            # Return minimal parsed data
            return {
                "skills": self._extract_skills_fallback(text),
                "parse_quality_score": 20,
            }

    def _extract_skills_fallback(self, text: str) -> List[str]:
        """Fallback skill extraction using pattern matching."""
        # Common tech skills to look for
        common_skills = [
            "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Rust",
            "React", "Angular", "Vue", "Node.js", "Django", "FastAPI", "Flask",
            "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform",
            "PostgreSQL", "MySQL", "MongoDB", "Redis", "Elasticsearch",
            "Git", "CI/CD", "Jenkins", "GitHub Actions",
            "Machine Learning", "Deep Learning", "NLP", "Computer Vision",
            "REST API", "GraphQL", "Microservices", "Agile", "Scrum",
        ]

        text_lower = text.lower()
        found_skills = []

        for skill in common_skills:
            if skill.lower() in text_lower:
                found_skills.append(skill)

        return found_skills

    def _apply_parsed_data(self, resume: Resume, data: Dict[str, Any]) -> None:
        """Apply parsed data to resume object."""
        # Basic info
        resume.full_name = data.get("full_name")
        resume.email = data.get("email")
        resume.phone = data.get("phone")
        resume.location = data.get("location")
        resume.linkedin_url = data.get("linkedin_url")
        resume.github_url = data.get("github_url")
        resume.portfolio_url = data.get("portfolio_url")
        resume.summary = data.get("summary")

        # Lists
        resume.skills = data.get("skills", [])
        resume.certifications = data.get("certifications", [])
        resume.languages = data.get("languages", [])

        # Sanitize education - ensure year and gpa are strings (LLM may return numbers)
        education = data.get("education", [])
        for edu in education:
            if edu.get("year") is not None:
                edu["year"] = str(edu["year"])
            if edu.get("gpa") is not None:
                edu["gpa"] = str(edu["gpa"])
        resume.education = education

        # Parse quality
        resume.parse_quality_score = data.get("parse_quality_score")

        # Work experiences
        for exp_data in data.get("work_experiences", []):
            exp = WorkExperience(
                resume=resume,
                company=exp_data.get("company", "Unknown"),
                title=exp_data.get("title", "Unknown"),
                location=exp_data.get("location"),
                employment_type=exp_data.get("employment_type"),
                is_remote=exp_data.get("is_remote", False),
                start_date=self._parse_date(exp_data.get("start_date")),
                end_date=self._parse_date(exp_data.get("end_date")),
                is_current=exp_data.get("is_current", False),
                description=exp_data.get("description"),
                achievements=exp_data.get("achievements", []),
                skills_used=exp_data.get("skills_used", []),
                metrics=exp_data.get("metrics", {}),
            )
            resume.work_experiences.append(exp)

    def _parse_date(self, date_str: Optional[str]) -> Optional[date]:
        """Parse date string to date object."""
        if not date_str:
            return None

        date_str = str(date_str).strip()

        # Try full ISO format
        try:
            return datetime.strptime(date_str, "%Y-%m-%d").date()
        except ValueError:
            pass

        # Try year-month
        try:
            return datetime.strptime(date_str, "%Y-%m").date()
        except ValueError:
            pass

        # Try just year
        try:
            year = int(date_str[:4])
            if 1950 <= year <= 2100:
                return date(year, 1, 1)
        except (ValueError, TypeError):
            pass

        return None


# Convenience functions
async def parse_resume_file(
    db: AsyncSession,
    user_id: UUID,
    file_content: bytes,
    file_name: str,
    file_type: str,
) -> Resume:
    """Parse a resume file for a user."""
    service = ResumeService(db)
    return await service.upload_and_parse(user_id, file_content, file_name, file_type)


async def get_user_resume(db: AsyncSession, user_id: UUID) -> Optional[Resume]:
    """Get a user's resume."""
    service = ResumeService(db)
    return await service.get_resume(user_id)
